import streamlit as st
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os
import time

# Function to retry Google Sheets API calls with exponential backoff
def safe_fetch_worksheet(worksheet, retries=5, delay=2):
    for attempt in range(retries):
        try:
            return worksheet.get_all_values()
        except Exception as e:
            st.warning(f"Retrying due to: {e}. Attempt {attempt + 1}")
            time.sleep(delay)
            delay *= 2
    return []

# Function to fetch and clean PRODUCTS sheet
def fetch_products_data(sheet):
    worksheet = sheet.worksheet("PRODUCTS")  # Access the PRODUCTS sheet
    raw_data = safe_fetch_worksheet(worksheet)
    if len(raw_data) > 1:
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
        # Process the columns
        df['Date Uploaded'] = pd.to_datetime(df['Date Uploaded:'], format='%d/%m/%Y', errors='coerce')
        df['eBay Price'] = pd.to_numeric(df['eBay Price'].str.replace(r'[^\d.]', '', regex=True), errors='coerce')
        df['Profit Per Product Present'] = pd.to_numeric(df['Profit Per Product Present'].str.replace('%', '', regex=True), errors='coerce') / 100
        return df[['Date Uploaded', 'eBay Price', 'Profit Per Product Present']]
    return pd.DataFrame()

# Function to fetch and clean ORDERS sheet
def fetch_orders_data(sheet):
    worksheet = sheet.worksheet("ORDERS")  # Access the ORDERS sheet
    raw_data = safe_fetch_worksheet(worksheet)
    if len(raw_data) > 1:
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
        # Process the columns
        df['Date Of Purchase'] = pd.to_datetime(df['Date Of Purchase'], format='%d/%m/%Y', errors='coerce')
        df['eBay Price'] = pd.to_numeric(df['eBay Price'].str.replace(r'[^\d.]', '', regex=True), errors='coerce')
        df['Profit Per Sale USD'] = pd.to_numeric(df['Profit Per Sale USD'].str.replace(r'[^\d.]', '', regex=True), errors='coerce')
        df['Total Sales'] = pd.to_numeric(df['Total Sales'].str.replace(r'[^\d.]', '', regex=True), errors='coerce')
        df['Total Profits'] = pd.to_numeric(df['Total Profits'].str.replace(r'[^\d.]', '', regex=True), errors='coerce')
        return df[['Date Of Purchase', 'eBay Price', 'Profit Per Sale USD', 'Total Sales', 'Total Profits']]
    return pd.DataFrame()

# Function to fetch and clean HOURSWORKED sheet
def fetch_hours_data(sheet):
    worksheet = sheet.worksheet("HOURSWORKED")  # Access the HOURSWORKED sheet
    raw_data = safe_fetch_worksheet(worksheet)
    if len(raw_data) > 1:
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
        # Process the columns
        df['Date'] = pd.to_datetime(df['Date:'], format='%d/%m/%Y', errors='coerce')
        df['Hours'] = pd.to_numeric(df['Hours:'], errors='coerce')
        return df[['Date', 'Hours']]
    return pd.DataFrame()

# Function to generate graphs
def create_graph(data, title, x_label, y_label, graph_type="bar"):
    if data.empty:
        return None
    data.index = data.index.astype(str)
    plt.figure(figsize=(6, 4))
    if graph_type == "bar":
        plt.bar(data.index, data.values)
    else:
        plt.plot(data.index, data.values, marker='o')
    plt.title(title)
    plt.xlabel(x_label)
    plt.ylabel(y_label)
    plt.xticks(rotation=45)
    buffer = BytesIO()
    plt.tight_layout()
    plt.savefig(buffer, format="png")
    buffer.seek(0)
    plt.close()
    return buffer

# Function to generate .docx report
def generate_report(client_name, products, orders, hours):
    doc = Document()
    doc.add_heading(f'Report for {client_name}', 0)

    # Extract totals
    total_sales = orders['Total Sales'].sum() if 'Total Sales' in orders else 0
    total_profit = orders['Total Profits'].sum() if 'Total Profits' in orders else 0
    total_hours = hours['Hours'].sum() if 'Hours' in hours else 0

    doc.add_paragraph(f"Total Sales: ${total_sales:.2f}")
    doc.add_paragraph(f"Total Profit: ${total_profit:.2f}")
    doc.add_paragraph(f"Total Hours Worked: {total_hours:.2f} hours")

    # Graphs
    doc.add_heading("Graphs", level=1)

    if not products.empty:
        products_per_month = products.groupby(products['Date Uploaded'].dt.to_period('M')).size()
        graph = create_graph(products_per_month, "Products Uploaded per Month", "Month", "Number of Products")
        if graph: doc.add_picture(graph, width=Inches(5))

    if not orders.empty:
        sales_per_month = orders.groupby(orders['Date Of Purchase'].dt.to_period('M'))['eBay Price'].count()
        graph = create_graph(sales_per_month, "Sales per Month", "Month", "Number of Sales")
        if graph: doc.add_picture(graph, width=Inches(5))

        profit_per_month = orders.groupby(orders['Date Of Purchase'].dt.to_period('M'))['Profit Per Sale USD'].sum()
        graph = create_graph(profit_per_month, "Profit per Month", "Month", "Profit ($)")
        if graph: doc.add_picture(graph, width=Inches(5))

    if not hours.empty:
        hours_per_month = hours.groupby(hours['Date'].dt.to_period('M'))['Hours'].sum()
        graph = create_graph(hours_per_month, "Hours Worked per Month", "Month", "Hours")
        if graph: doc.add_picture(graph, width=Inches(5))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App
st.title("Client Report Generator")

try:
    credentials_path = "credentials.json"
    if not os.path.exists(credentials_path):
        st.error("The `credentials.json` file is missing.")
    else:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_name(credentials_path, scope)
        client = gspread.authorize(creds)

        sheet_links_input = st.text_area("Paste Google Sheets Links (one per line):", key="google_sheets_input")

        if st.button("Generate Reports"):
            sheet_links = [link.strip() for link in sheet_links_input.split("\n") if link.strip()]
            for idx, sheet_url in enumerate(sheet_links):
                try:
                    st.write(f"Processing link {idx + 1}...")
                    client_sheet = client.open_by_url(sheet_url)
                    client_name = client_sheet.title

                    products = fetch_products_data(client_sheet)
                    orders = fetch_orders_data(client_sheet)
                    hours = fetch_hours_data(client_sheet)

                    report_buffer = generate_report(client_name, products, orders, hours)

                    st.download_button(
                        label=f"Download {client_name}'s Report",
                        data=report_buffer,
                        file_name=f"{client_name}_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Error processing link {sheet_url}: {e}")
except Exception as e:
    st.error(f"An error occurred: {e}")
